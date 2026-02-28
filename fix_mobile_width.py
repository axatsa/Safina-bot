from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

def fix_table_width_for_mobile():
    template_path = "backend/app/services/docx/template.docx"
    doc = Document(template_path)
    
    # We know the items table is doc.tables[2] based on previous logs
    t = doc.tables[2]
    
    # 1. Set table alignment to center (optional but good for mobile)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 2. Force table width to 100% (autofit to window usually handles this)
    t.autofit = True
    
    # A more aggressive XML-level fix for 100% width:
    # Set the w:tblW to type 'pct' (percent) and w:w to '5000' (which means 100% in Word OOXML)
    tbl_pr = t._tbl.tblPr
    if tbl_pr is not None:
        tbl_w = tbl_pr.xpath('./w:tblW')
        if tbl_w:
            tbl_w[0].set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'pct')
            tbl_w[0].set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '5000')

    doc.save(template_path)
    print("Table width set to 100% for mobile compatibility.")

if __name__ == "__main__":
    fix_table_width_for_mobile()
