from docx import Document
import os

# Script to precisely fix the docxtpl tags in the template for row repetition.
# This ensures that the tags are not split by Word's internal XML formatting.

def fix_template():
    template_path = "backend/app/services/docx/template.docx"
    if not os.path.exists(template_path):
        print(f"Template not found at {template_path}")
        return

    doc = Document(template_path)
    
    # 1. Find the table with our item tags
    target_table = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "item.name" in cell.text or "item" in cell.text:
                    target_table = table
                    break
            if target_table: break
        if target_table: break

    if not target_table:
        print("Could not find the items table.")
        return

    # 2. Identify the data row (usually the one after the header)
    # Looking for a row that has item-related placeholders
    data_row_idx = -1
    for i, row in enumerate(target_table.rows):
        if any("item" in cell.text for cell in row.cells):
            data_row_idx = i
            break

    if data_row_idx == -1:
        print("Could not find the data row in the table.")
        return

    print(f"Found data row at index {data_row_idx}. Fixing tags...")

    # 3. Precisely set the tags in the data row cells
    row = target_table.rows[data_row_idx]
    
    # Cell 0: No (Loop Start)
    cell_no = row.cells[0]
    cell_no.text = "" # Clear
    p = cell_no.paragraphs[0]
    p.add_run("{% for item in items %}").bold = False
    p.add_run("{{ item.no }}")

    # Cell 1: Name
    cell_name = row.cells[1]
    cell_name.text = "{{ item.name }}"

    # Cell 2: Quantity
    cell_qty = row.cells[2]
    cell_qty.text = "{{ item.quantity }}"

    # Cell 3: Price/Total (Loop End)
    # Based on the user's template, they have Price and then Summa.
    # Let's assume the last cell is the one for loop end.
    last_cell_idx = len(row.cells) - 1
    cell_total = row.cells[last_cell_idx]
    cell_total.text = ""
    p = cell_total.paragraphs[0]
    p.add_run("{{ item.price }}") # User used item.price in the last test
    p.add_run("{% endfor %}")

    # 4. Save the fixed template
    doc.save(template_path)
    print("Template fixed and saved successfully!")

if __name__ == "__main__":
    fix_template()
